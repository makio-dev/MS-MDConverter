"""Markdown to Excel/Word Converter"""

import sys
import os
import re
from pathlib import Path

import openpyxl
from openpyxl.styles import Font, Alignment, PatternFill, Border, Side
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


# ─── Markdown Parser ────────────────────────────────────────────────────────

class MDBlock:
    """Parsed markdown block."""
    def __init__(self, kind, content, level=0):
        self.kind = kind        # "heading", "table", "paragraph", "list", "code", "hr"
        self.content = content  # str or list (table rows)
        self.level = level      # heading level (1-6) or list nesting


def parse_md(text: str) -> list[MDBlock]:
    """Parse markdown text into a list of blocks."""
    lines = text.split("\n")
    blocks: list[MDBlock] = []
    i = 0

    while i < len(lines):
        line = lines[i]

        # blank line
        if not line.strip():
            i += 1
            continue

        # heading (ATX style)
        m = re.match(r"^(#{1,6})\s+(.+)$", line)
        if m:
            level = len(m.group(1))
            blocks.append(MDBlock("heading", m.group(2).strip(), level))
            i += 1
            continue

        # horizontal rule
        if re.match(r"^(\*{3,}|-{3,}|_{3,})\s*$", line):
            blocks.append(MDBlock("hr", ""))
            i += 1
            continue

        # fenced code block
        if line.strip().startswith("```"):
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith("```"):
                code_lines.append(lines[i])
                i += 1
            i += 1  # skip closing ```
            blocks.append(MDBlock("code", "\n".join(code_lines)))
            continue

        # table
        if "|" in line and i + 1 < len(lines) and re.match(r"^\|?[\s\-:|]+\|", lines[i + 1]):
            table_rows = []
            # header row
            table_rows.append(_parse_table_row(line))
            i += 1  # skip separator line
            i += 1
            while i < len(lines) and "|" in lines[i] and lines[i].strip():
                table_rows.append(_parse_table_row(lines[i]))
                i += 1
            blocks.append(MDBlock("table", table_rows))
            continue

        # unordered list
        if re.match(r"^(\s*)([-*+])\s+", line):
            list_items = []
            while i < len(lines) and re.match(r"^(\s*)([-*+])\s+", lines[i]):
                m = re.match(r"^(\s*)([-*+])\s+(.+)$", lines[i])
                indent = len(m.group(1))
                list_items.append((indent, m.group(3).strip()))
                i += 1
            blocks.append(MDBlock("list", list_items))
            continue

        # ordered list
        if re.match(r"^(\s*)\d+\.\s+", line):
            list_items = []
            while i < len(lines) and re.match(r"^(\s*)\d+\.\s+", lines[i]):
                m = re.match(r"^(\s*)\d+\.\s+(.+)$", lines[i])
                indent = len(m.group(1))
                list_items.append((indent, m.group(2).strip()))
                i += 1
            blocks.append(MDBlock("list", list_items))
            continue

        # paragraph (collect consecutive non-empty lines)
        para_lines = []
        while i < len(lines) and lines[i].strip() and not _is_special_line(lines[i], lines, i):
            para_lines.append(lines[i].strip())
            i += 1
        if para_lines:
            blocks.append(MDBlock("paragraph", " ".join(para_lines)))

    return blocks


def _parse_table_row(line: str) -> list[str]:
    """Parse a markdown table row into cells."""
    line = line.strip()
    if line.startswith("|"):
        line = line[1:]
    if line.endswith("|"):
        line = line[:-1]
    return [cell.strip() for cell in line.split("|")]


def _is_special_line(line: str, lines: list[str], i: int) -> bool:
    """Check if a line starts a new special block."""
    if re.match(r"^#{1,6}\s+", line):
        return True
    if re.match(r"^(\*{3,}|-{3,}|_{3,})\s*$", line):
        return True
    if line.strip().startswith("```"):
        return True
    if "|" in line and i + 1 < len(lines) and re.match(r"^\|?[\s\-:|]+\|", lines[i + 1]):
        return True
    if re.match(r"^(\s*)([-*+])\s+", line):
        return True
    if re.match(r"^(\s*)\d+\.\s+", line):
        return True
    return False


def strip_inline_md(text: str) -> str:
    """Remove inline markdown formatting (bold, italic, code, links).
    Converts <br> / <br/> / <br /> tags to newlines."""
    text = re.sub(r"<br\s*/?>", "\n", text, flags=re.IGNORECASE)
    text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
    text = re.sub(r"__(.+?)__", r"\1", text)
    text = re.sub(r"\*(.+?)\*", r"\1", text)
    text = re.sub(r"_(.+?)_", r"\1", text)
    text = re.sub(r"`(.+?)`", r"\1", text)
    text = re.sub(r"\[(.+?)\]\(.+?\)", r"\1", text)
    text = re.sub(r'<a\s[^>]*?>(.*?)</a>', r"\1", text, flags=re.IGNORECASE)
    return text


# ─── Excel Exporter ─────────────────────────────────────────────────────────

HEADING_FILLS = {
    1: PatternFill(start_color="1F4E79", end_color="1F4E79", fill_type="solid"),
    2: PatternFill(start_color="2E75B6", end_color="2E75B6", fill_type="solid"),
    3: PatternFill(start_color="9DC3E6", end_color="9DC3E6", fill_type="solid"),
}

HEADING_FONTS = {
    1: Font(name="Yu Gothic", size=16, bold=True, color="FFFFFF"),
    2: Font(name="Yu Gothic", size=14, bold=True, color="FFFFFF"),
    3: Font(name="Yu Gothic", size=12, bold=True, color="1F4E79"),
    4: Font(name="Yu Gothic", size=11, bold=True),
    5: Font(name="Yu Gothic", size=10, bold=True),
    6: Font(name="Yu Gothic", size=10, bold=True, italic=True),
}

THIN_BORDER = Border(
    left=Side(style="thin"),
    right=Side(style="thin"),
    top=Side(style="thin"),
    bottom=Side(style="thin"),
)

TABLE_HEADER_FILL = PatternFill(start_color="D6E4F0", end_color="D6E4F0", fill_type="solid")
TABLE_HEADER_FONT = Font(name="Yu Gothic", size=10, bold=True)
TABLE_BODY_FONT = Font(name="Yu Gothic", size=10)
BODY_FONT = Font(name="Yu Gothic", size=10)
CODE_FONT = Font(name="Consolas", size=9)
CODE_FILL = PatternFill(start_color="F2F2F2", end_color="F2F2F2", fill_type="solid")


def export_to_excel(blocks: list[MDBlock], output_path: str):
    """Export parsed markdown blocks to an Excel file."""
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Markdown"
    ws.sheet_properties.outlinePr = openpyxl.worksheet.properties.Outline(
        summaryBelow=False, summaryRight=False
    )

    row = 1

    for block in blocks:
        if block.kind == "heading":
            level = block.level
            text = strip_inline_md(block.content)
            cell = ws.cell(row=row, column=1, value=text)
            cell.font = HEADING_FONTS.get(level, HEADING_FONTS[6])
            if level in HEADING_FILLS:
                for col in range(1, 8):
                    ws.cell(row=row, column=col).fill = HEADING_FILLS[level]
            cell.alignment = Alignment(vertical="center")
            ws.row_dimensions[row].height = 28 if level <= 2 else 22
            row += 1

        elif block.kind == "table":
            table_rows = block.content
            if not table_rows:
                continue
            num_cols = max(len(r) for r in table_rows)
            for ri, trow in enumerate(table_rows):
                for ci, cell_text in enumerate(trow):
                    cell = ws.cell(row=row, column=ci + 1, value=strip_inline_md(cell_text))
                    cell.border = THIN_BORDER
                    cell.alignment = Alignment(wrap_text=True, vertical="center")
                    if ri == 0:
                        cell.fill = TABLE_HEADER_FILL
                        cell.font = TABLE_HEADER_FONT
                    else:
                        cell.font = TABLE_BODY_FONT
                row += 1
            row += 1  # blank row after table

        elif block.kind == "list":
            for indent, text in block.content:
                prefix = "  " * (indent // 2) + "• "
                cell = ws.cell(row=row, column=1, value=prefix + strip_inline_md(text))
                cell.font = BODY_FONT
                cell.alignment = Alignment(wrap_text=False)
                row += 1
            row += 1

        elif block.kind == "code":
            for code_line in block.content.split("\n"):
                cell = ws.cell(row=row, column=1, value=code_line)
                cell.font = CODE_FONT
                cell.fill = CODE_FILL
                cell.alignment = Alignment(wrap_text=False)
                row += 1
            row += 1

        elif block.kind == "hr":
            for col in range(1, 8):
                ws.cell(row=row, column=col).border = Border(
                    bottom=Side(style="medium", color="888888")
                )
            row += 1

        elif block.kind == "paragraph":
            cell = ws.cell(row=row, column=1, value=strip_inline_md(block.content))
            cell.font = BODY_FONT
            cell.alignment = Alignment(wrap_text=False)
            row += 1

    # auto-adjust column widths
    for col_cells in ws.columns:
        max_len = 0
        col_letter = col_cells[0].column_letter
        for cell in col_cells:
            if cell.value:
                max_len = max(max_len, len(str(cell.value)))
        ws.column_dimensions[col_letter].width = min(max_len + 4, 60)

    wb.save(output_path)
    print(f"Excel exported: {output_path}")


# ─── Word Exporter ──────────────────────────────────────────────────────────

def export_to_word(blocks: list[MDBlock], output_path: str):
    """Export parsed markdown blocks to a Word file."""
    doc = Document()

    # set default font
    style = doc.styles["Normal"]
    style.font.name = "Yu Gothic"
    style.font.size = Pt(10)

    for block in blocks:
        if block.kind == "heading":
            level = min(block.level, 4)  # Word supports Heading 1-4 easily
            text = strip_inline_md(block.content)
            doc.add_heading(text, level=level)

        elif block.kind == "table":
            table_rows = block.content
            if not table_rows:
                continue
            num_cols = max(len(r) for r in table_rows)
            table = doc.add_table(rows=len(table_rows), cols=num_cols)
            table.style = "Table Grid"
            # auto-fit table to window width
            table.autofit = True
            from docx.oxml.ns import qn
            tbl = table._tbl
            tbl_pr = tbl.tblPr if tbl.tblPr is not None else tbl._add_tblPr()
            from lxml import etree
            # Remove fixed width, use auto layout
            for existing in tbl_pr.findall(qn("w:tblW")):
                tbl_pr.remove(existing)
            etree.SubElement(tbl_pr, qn("w:tblW"), attrib={
                qn("w:w"): "5000", qn("w:type"): "pct",
            })
            for ri, trow in enumerate(table_rows):
                for ci, cell_text in enumerate(trow):
                    cell = table.cell(ri, ci)
                    text = strip_inline_md(cell_text)
                    # Handle newlines (from <br>) as separate paragraphs in cell
                    lines = text.split("\n")
                    cell.text = lines[0]
                    for extra_line in lines[1:]:
                        cell.add_paragraph(extra_line)
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.size = Pt(9)
                    if ri == 0:
                        tc_pr = cell._element.get_or_add_tcPr()
                        shading_elm = etree.SubElement(
                            tc_pr, qn("w:shd"),
                            attrib={
                                qn("w:val"): "clear",
                                qn("w:color"): "auto",
                                qn("w:fill"): "D6E4F0",
                            },
                        )
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.bold = True
            doc.add_paragraph()  # spacing after table

        elif block.kind == "list":
            for indent, text in block.content:
                p = doc.add_paragraph(strip_inline_md(text), style="List Bullet")
                p.paragraph_format.left_indent = Inches(0.25 * (indent // 2))

        elif block.kind == "code":
            p = doc.add_paragraph()
            run = p.add_run(block.content)
            run.font.name = "Consolas"
            run.font.size = Pt(9)
            # light gray background via shading
            from docx.oxml.ns import qn
            from lxml import etree
            shd = etree.SubElement(
                p._element.get_or_add_pPr(), qn("w:shd"),
                attrib={
                    qn("w:val"): "clear",
                    qn("w:color"): "auto",
                    qn("w:fill"): "F2F2F2",
                },
            )

        elif block.kind == "hr":
            p = doc.add_paragraph()
            p_fmt = p.paragraph_format
            from docx.oxml.ns import qn
            from lxml import etree
            pPr = p._element.get_or_add_pPr()
            pBdr = etree.SubElement(pPr, qn("w:pBdr"))
            etree.SubElement(
                pBdr, qn("w:bottom"),
                attrib={
                    qn("w:val"): "single",
                    qn("w:sz"): "6",
                    qn("w:space"): "1",
                    qn("w:color"): "888888",
                },
            )

        elif block.kind == "paragraph":
            text = strip_inline_md(block.content)
            lines = text.split("\n")
            p = doc.add_paragraph(lines[0])
            for extra_line in lines[1:]:
                p.add_run("\n" + extra_line)

    doc.save(output_path)
    print(f"Word exported: {output_path}")


# ─── Main ───────────────────────────────────────────────────────────────────

def main():
    print("=" * 50)
    print("  Markdown Converter (MD → Excel / Word)")
    print("=" * 50)
    print()

    # ask for output format
    print("出力形式を選択してください:")
    print("  [1] Excel (.xlsx)")
    print("  [2] Word  (.docx)")
    print()

    while True:
        choice = input("番号を入力 (1 or 2): ").strip()
        if choice in ("1", "2"):
            break
        print("1 または 2 を入力してください。")

    fmt = "excel" if choice == "1" else "word"
    ext = ".xlsx" if fmt == "excel" else ".docx"

    # ask for input file
    print()
    md_path = input("変換するMarkdownファイルのパスを入力: ").strip().strip('"').strip("'")

    if not os.path.isfile(md_path):
        print(f"エラー: ファイルが見つかりません: {md_path}")
        sys.exit(1)

    # read markdown
    with open(md_path, encoding="utf-8") as f:
        md_text = f.read()

    # parse
    blocks = parse_md(md_text)

    if not blocks:
        print("エラー: Markdownの内容が空です。")
        sys.exit(1)

    # output path
    base = Path(md_path).stem
    output_dir = Path(md_path).parent
    output_path = str(output_dir / (base + ext))

    # export
    if fmt == "excel":
        export_to_excel(blocks, output_path)
    else:
        export_to_word(blocks, output_path)

    print()
    print("変換が完了しました！")


if __name__ == "__main__":
    main()
